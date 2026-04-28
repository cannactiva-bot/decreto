"""Aplica INSERCIÓN 14 (autocrítica + cannabinoides sintéticos + inversión art. 34 TFUE)
sobre v4.docx, preservando todas las modificaciones manuales del usuario y las 13 inserciones previas.

Acciones:
1. Inserta nueva subsección §4.5 'Una autocrítica necesaria…' tras la última frase de §4.4.
2. Renombra el actual §4.5 'Qué pedimos para este frente' a §4.6 (mismo contenido).
3. Reformula la petición (b) del §4.6 con cierre conectado a §4.5.
4. Añade nuevos documentos al §9 (Orden SND/380/2025, EMCDDA HHC 2023, PNSD Dossier 2025, OEDA SEAT 2025).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'
DST = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'

doc = Document(SRC)

_pStyle_cache = {}

def get_pStyle_xml(style_name):
    if style_name in _pStyle_cache:
        cached = _pStyle_cache[style_name]
        return deepcopy(cached) if cached is not None else None
    pStyle_el = None
    for p in doc.paragraphs:
        try:
            if p.style and p.style.name == style_name:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    found = pPr.find(qn('w:pStyle'))
                    if found is not None:
                        pStyle_el = found
                        break
        except Exception:
            continue
    _pStyle_cache[style_name] = pStyle_el
    return deepcopy(pStyle_el) if pStyle_el is not None else None


def apply_style_xml(p, style_name):
    pStyle_new = get_pStyle_xml(style_name)
    if pStyle_new is None:
        try:
            p.style = style_name
        except KeyError:
            pass
        return
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    existing = pPr.find(qn('w:pStyle'))
    if existing is not None:
        pPr.remove(existing)
    pPr.insert(0, pStyle_new)


def make_run(p, text, italic=False, bold=False):
    r = p.add_run(text)
    if italic:
        r.italic = True
    if bold:
        r.font.bold = True
    return r


def insert_after(target_p, text='', style=None, italic=False, bold=False, indent_quote=False):
    new_p = doc.add_paragraph()
    if style:
        apply_style_xml(new_p, style)
    if text:
        make_run(new_p, text, italic=italic, bold=bold)
    if indent_quote:
        new_p.paragraph_format.left_indent = Inches(0.4)
    target_p._element.addnext(new_p._element)
    return new_p


def insert_chain(target_p, paragraphs):
    cur = target_p
    for spec in paragraphs:
        if isinstance(spec, str):
            text, opts = spec, {}
        else:
            text, opts = spec
        cur = insert_after(cur, text, **opts)
    return cur


def find_para(text_substring):
    for p in doc.paragraphs:
        if text_substring in p.text:
            return p
    return None


def replace_text(p, old, new):
    full = p.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    runs = p.runs
    if not runs:
        p.add_run(new_full)
        return True
    for r in runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ''
    return True


# ============================================================
# INSERCIÓN 14 · §4.5 nueva (autocrítica + HHC + inversión art. 34)
# ============================================================
# Anchor: último párrafo del §4.4 (cierre del bloque sobre el Recurso TS)
p_44_close = find_para("ambas perspectivas son complementarias y no excluyentes")

ins14_paras = [
    # ===== Heading 2 nuevo §4.5 =====
    ("4.5 Una autocrítica necesaria · El punto ciego del cultivo y el riesgo de los cannabinoides sintéticos",
     {'style': 'Heading 2'}),

    # ===== §4.5.1 =====
    ("4.5.1 La autocrítica del propio sector", {'style': 'Heading 3'}),
    "Antes de formular nuestras peticiones a la Comisión Nacional de los Mercados y la Competencia, AECANI desea explicitar una autocrítica sectorial que considera necesaria por coherencia y por eficacia argumental: la liberación regulatoria del cáñamo industrial no exige la ausencia de control. El cáñamo industrial es un cultivo agrícola legítimo, pero es también —por su parecido morfológico con el cannabis psicoactivo y por la continuidad química entre sus principios activos y los de las variedades estupefacientes— un cultivo con sensibilidades específicas. La industria que AECANI representa no pide ausencia de control; pide control en el lugar correcto de la cadena de valor.",

    # ===== §4.5.2 =====
    ("4.5.2 Localización del punto ciego", {'style': 'Heading 3'}),
    "El sistema regulatorio español sobre el cáñamo industrial presenta hoy una asimetría: los productos finales están razonablemente cubiertos —el cannabis medicinal por el RD 903/2025; los productos a base de hierbas para fumar por la Directiva 2014/40/UE TPD y el RD 579/2017; los cosméticos por el Reglamento (CE) 1223/2009; los productos alimentarios por el régimen Novel Food del Reglamento (UE) 2015/2283—. En cambio, los eslabones de cultivo y de primera transformación se encuentran en un punto ciego regulatorio. No existe un régimen específico de control del cultivo del cáñamo industrial análogo al que España aplica al tabaco (RD 969/2014), al vino (Ley 24/2003), al lúpulo o a la adormidera. Tampoco existe un régimen de trazabilidad obligatoria de la primera transformación de la flor.",

    # ===== §4.5.3 =====
    ("4.5.3 El riesgo concreto: HHC y cannabinoides semi-sintéticos", {'style': 'Heading 3'}),
    "Este punto ciego no es teórico. Su manifestación más documentada es el desplazamiento del CBD obtenido del cáñamo industrial hacia la fabricación clandestina de HHC (hexahidrocannabinol) y otros cannabinoides semi-sintéticos (HHC-O, THCP, THC-O, H4-CBD y derivados análogos), que se sintetizan químicamente a partir del CBD aislado de la flor de cáñamo industrial.",
    "El European Monitoring Centre for Drugs and Drug Addiction (hoy European Union Drugs Agency, EUDA) describió con precisión el mecanismo en su informe técnico de mayo de 2023 [Anexo 9.13]:",
    ("«The current large-scale manufacturing of HHC is based on low-THC hemp derived CBD-extract which is first transformed into a mixture of Δ8-THC and Δ9-THC followed by catalytic hydrogenation of the THC isomer mixture into the final product.»",
     {'italic': True, 'indent_quote': True}),
    "Y, sobre la causa estructural:",
    ("«This new market is linked to: the legalisation of hemp cultivation in the US in 2018; subsequent abundant / surplus supply of hemp and cannabidiol (CBD) derived from hemp that can be used as a precursor for SSCs [semi-synthetic cannabinoids].»",
     {'italic': True, 'indent_quote': True}),
    "El propio Plan Nacional sobre Drogas (PNSD), en su Dossier sobre Cannabinoides Sintéticos de 11 de abril de 2025 [Anexo 9.14], reconoce explícitamente el modus operandi del mercado ilícito español:",
    ("«Para maximizar los beneficios, hay traficantes que impregnan con cannabinoides sintéticos el cáñamo industrial con bajo contenido de THC, que posee un aspecto similar al de la hierba de cannabis.»",
     {'italic': True, 'indent_quote': True}),
    "Las cifras dimensionan el problema. La EUDA vigila a 8 de octubre de 2025 un total de 298 cannabinoides; en 2024 se notificaron 20 nuevos cannabinoides (18 de ellos semi-sintéticos), más del 40% de las nuevas sustancias psicoactivas notificadas a la EUDA ese año. En España, el Sistema Español de Alerta Temprana (SEAT) registró 28 intoxicaciones agudas en servicios de urgencias en 2024, en su mayoría asociadas al consumo de gominolas y caramelos adulterados con HHC y THCP, con 14 notificaciones concentradas en Madrid y Barcelona [Informe SEAT-OEDA Resumen Ejecutivo 2025, Anexo 9.15].",

    # ===== §4.5.4 =====
    ("4.5.4 La asimetría regulatoria española", {'style': 'Heading 3'}),
    "España ha adoptado una respuesta regulatoria downstream. La Orden SND/380/2025, de 14 de abril (BOE 22-abr-2025) [Anexo 5.11] incluye HHC, HHC-O, HHCP, HHCP-O, THCP, delta-8-THCP, THC-O, THCP-O, H4-CBD y delta-9-THCA en la Lista II anexa al RD 2829/1977, con entrada en vigor el 23 de abril de 2025. Es una respuesta sobre la molécula final que cabe valorar como necesaria.",
    "Pero la Orden SND/380/2025 es insuficiente desde la perspectiva regulatoria. La causa estructural identificada por la EMCDDA es la «abundant / surplus supply of hemp and CBD derived from hemp»: un fenómeno upstream que ninguna fiscalización de la molécula final, por sí sola, puede prevenir. España carece a día de hoy de un régimen específico de control del cultivo y de la primera transformación del cáñamo industrial que actúe sobre la causa estructural del problema. La asimetría es completa: la administración fiscaliza la molécula sintética producida por el desvío después de que el desvío se haya producido, mientras que la materia prima que alimenta ese desvío —el CBD extraído de la flor— no está sujeta a régimen de trazabilidad alguno.",
    "Simultáneamente, la pretensión regulatoria contenida en el TRIS 2025/0044/ES no aborda este riesgo: somete a presión adicional a los operadores legítimos del cáñamo industrial —los registrados, los que tributan, los que mantienen los 1.700 establecimientos minoristas y los 6.700 empleos documentados en el §3.4— sin añadir control alguno sobre la rama del sector que efectivamente se desvía hacia la fabricación ilícita de cannabinoides semi-sintéticos.",

    # ===== §4.5.5 — la inversión art. 34 TFUE =====
    ("4.5.5 La inversión del argumento de salud pública: la inacción como riesgo", {'style': 'Heading 3'}),
    "El artículo 34 del Tratado de Funcionamiento de la Unión Europea prohíbe a los Estados miembros toda restricción cuantitativa a la libre circulación intracomunitaria de mercancías y toda medida de efecto equivalente. El artículo 36 del mismo Tratado admite excepciones por motivos de orden público, salud pública u otros taxativamente enumerados, pero la jurisprudencia del Tribunal de Justicia de la Unión Europea ha establecido —en términos especialmente claros en la sentencia C-663/18 Kanavape, ya citada— que esas excepciones requieren la acreditación, por parte del Estado que las invoque, de un riesgo real para la salud pública apreciado a la luz de los datos científicos más recientes, y que su aplicación debe respetar el principio de proporcionalidad.",
    "La administración española, en relación con el cáñamo industrial y sus productos derivados, opera hoy en sentido inverso al estándar exigido por el Derecho de la Unión: invoca implícitamente la salud pública como justificación de un régimen restrictivo basado en definiciones expansivas (asimilación de toda materia prima vegetal de cannabis al régimen de los estupefacientes, con independencia de su contenido en THC y de su quimiotipo) y, en su caso, en consideraciones de orden público (ausencia de un sistema unificado de verificación de edad, ausencia de canal regulado para el consumidor adulto), sin acreditar el riesgo real exigido por la jurisprudencia Kanavape ni adoptar las medidas de regulación positiva que harían tales consideraciones razonables.",
    "El resultado es paradójico y debe explicitarse con claridad: la administración española invoca la salud pública para impedir el comercio del cáñamo industrial legítimo, pero su inacción regulatoria simultáneamente genera los riesgos para la salud pública que dice querer prevenir. La constatación es directa:",
    ("Operadores no trazados. El comercio minorista del cáñamo industrial (≈1.700 establecimientos, ≈450 M€ de facturación, ≈6.700 empleos directos) carece de código CNAE específico, de epígrafe IAE propio y de cualquier sistema de registro o control administrativo diferenciado. La administración no sabe quién opera en el sector. Esto, per se, es un riesgo regulatorio.",
     {'style': 'List Paragraph'}),
    ("Cultivo no trazado. España subvenciona vía PAC el cultivo del cáñamo industrial pero no aplica trazabilidad obligatoria ni régimen de contratación productor-transformador. La materia prima destinada a la extracción de CBD —el precursor químico de los cannabinoides semi-sintéticos identificado por la EMCDDA— circula sin control diferenciado.",
     {'style': 'List Paragraph'}),
    ("Producto en mercado gris. La administración no admite la notificación de productos a base de hierbas para fumar de cáñamo industrial bajo el régimen general de la Directiva 2014/40/UE TPD (cf. §5.1 del presente dossier), pero los mismos productos se siguen comercializando bajo otras categorías (referencia DGT V2242-22 sobre flores ornamentales tributadas al 10% IVA). El consumidor adulto adquiere el producto sin garantías de procedencia ni de composición.",
     {'style': 'List Paragraph'}),
    ("Desvío al mercado ilícito. En este vacío proliferan los cannabinoides semi-sintéticos descritos en §4.5.3, con manifestación documentada en intoxicaciones agudas en servicios de urgencias españoles (28 casos notificados en 2024, principalmente por gominolas adulteradas con HHC y THCP).",
     {'style': 'List Paragraph'}),
    "El verdadero riesgo para la salud pública que la administración invoca como justificación de su política restrictiva no procede del cáñamo industrial en sí —cuyo perfil de riesgo está cualificado por la OMS (ECDD 2018), por la EMCDDA y por la jurisprudencia Kanavape como bajo cuando se trata del producto con bajo contenido en THC— sino de la inacción regulatoria de la propia administración, que simultáneamente impide la operación legal del sector y se desentiende del control efectivo del riesgo aguas arriba. Dicho con la simetría que el caso reclama: en el estado actual, no tener trazados a los operadores que comercializan flor de cáñamo (sea como producto ornamental bajo la ficción tributaria del 10% IVA, sea como producto a base de hierbas para fumar bajo régimen general TPD del que la administración los excluye) es objetivamente peor para la salud pública que regularlos correctamente.",
    "La inversión argumental es, por tanto, simétrica y plenamente proyectable a la propia CNMC en sede de competencia: la inacción regulatoria no solo produce distorsiones de competencia (las documentadas en el §5 del presente dossier), sino que produce además los efectos sanitarios que la propia administración invoca para perpetuar la inacción. La salida razonable a este círculo es exactamente la que AECANI propone: regular —de manera diseñada, proporcionada, coherente con el Derecho de la Unión y con los modelos espejo nacionales (tabaco, vino, lúpulo, adormidera)— en lugar de prohibir.",

    # ===== §4.5.6 conclusion =====
    ("4.5.6 La conclusión: control aguas arriba, no presión aguas abajo", {'style': 'Heading 3'}),
    "La consecuencia de política regulatoria es directa. La forma eficaz de prevenir el desvío del cáñamo industrial a la fabricación de HHC y cannabinoides semi-sintéticos no consiste en restringir el comercio minorista del cáñamo industrial al consumidor adulto; consiste en establecer un régimen específico de control de cultivos y de primera transformación, con trazabilidad desde la semilla hasta el primer transformador autorizado, contratación obligatoria entre productor y transformador, y registro estadístico diferenciado. Es el modelo que España aplica con éxito al tabaco, al vino, al lúpulo y a la adormidera; es el modelo que la propia Propuesta COM(2025) 553 final invita implícitamente al reconocer todas las partes de la planta de cáñamo (incluidas las sumidades floridas bajo NC 1211 90 86) como producto agrícola comercializable dentro de la Organización Común de Mercados; y es el modelo que la EMCDDA identifica como condición para gobernar el riesgo de los cannabinoides semi-sintéticos en su origen.",
    "Un régimen de trazabilidad correctamente diseñado reduciría a niveles residuales la posibilidad de que la flor de cáñamo industrial cultivada legalmente en España se desvíe a la fabricación clandestina de HHC u otros cannabinoides semi-sintéticos. La prevención de este riesgo requiere actuación regulatoria, sí; pero requiere actuación en el lugar adecuado de la cadena de valor —el cultivo y la primera transformación— y no en el comerciante final del producto legítimo, que es el único eslabón actualmente sujeto a presión regulatoria por parte de la administración española.",

    # ===== §4.5.7 conexion con peticiones =====
    ("4.5.7 Conexión con las peticiones", {'style': 'Heading 3'}),
    "Esta autocrítica no debilita las peticiones que AECANI formula a continuación; las refuerza. Solicitar a la CNMC un pronunciamiento técnico sobre el artículo 5.2 del RD 903/2025 y una recomendación al Gobierno para la elaboración de una norma específica de control de cultivos no es una pretensión liberalizadora: es una pretensión de regulación bien diseñada. El sector que representa AECANI es el primer interesado en que esa regulación exista, porque es la condición para que la rama legítima del cáñamo industrial pueda operar con seguridad jurídica y la rama ilegítima —la del desvío a HHC y a cannabinoides semi-sintéticos— desaparezca por sustitución.",
]

if p_44_close:
    insert_chain(p_44_close, ins14_paras)
    print('INSERCIÓN 14 (§4.5 nueva): OK')
else:
    print('INSERCIÓN 14: FAIL — anchor "ambas perspectivas son complementarias" no encontrado')


# ============================================================
# Renombrar §4.5 antiguo a §4.6
# ============================================================
p_45_old = find_para("4.5 Qué pedimos para este frente")
if p_45_old:
    ok = replace_text(p_45_old, "4.5 Qué pedimos para este frente", "4.6 Qué pedimos para este frente")
    print(f'Renumeración §4.5→§4.6: {"OK" if ok else "FAIL"}')

# Reformular petición (b) del §4.6 (antes §4.5) — añadir cierre conectado a §4.5
p_pet_b = find_para("Recomendación al Gobierno para la elaboración de una norma específica de control de cultivos")
if p_pet_b:
    old_close = "con materia prima vendible sin licencia de estupefacientes a procesadores y a la industria farmacéutica autorizada."
    new_close = "con materia prima vendible sin licencia de estupefacientes a procesadores y a la industria farmacéutica autorizada, y con la finalidad expresa, conforme al §4.5 anterior, de cerrar el punto ciego regulatorio identificado y prevenir el desvío del cáñamo industrial a la fabricación clandestina de cannabinoides semi-sintéticos (HHC, HHC-O, THCP y derivados análogos)."
    ok = replace_text(p_pet_b, old_close, new_close)
    print(f'Reformulación petición (b) §4.6: {"OK" if ok else "FAIL — ya pudo aplicarse texto distinto"}')


# ============================================================
# Adición §9 — referencias HHC y cannabinoides sintéticos
# ============================================================
# Insertar nuevo subbloque al final de §9 (antes de "Fin del dossier")
p_fin = find_para("Fin del dossier")
if p_fin is None:
    p_fin = doc.paragraphs[-1]

ins_ref_hhc = [
    ("Cannabinoides sintéticos y HHC (cf. §4.5)", {'style': 'Heading 2'}),
    ("EMCDDA · Hexahydrocannabinol (HHC) and related substances · Technical report, mayo 2023.",
     {'style': 'List Paragraph'}),
    ("Plan Nacional sobre Drogas · Dossier «Cannabinoides Sintéticos», 11 de abril de 2025.",
     {'style': 'List Paragraph'}),
    ("OEDA · Informe SEAT 2025 (Resumen Ejecutivo) · Sistema Español de Alerta Temprana sobre Nuevas Sustancias Psicoactivas.",
     {'style': 'List Paragraph'}),
    ("Orden SND/380/2025, de 14 de abril, por la que se incluyen HHC y otros cannabinoides semi-sintéticos en la Lista II anexa al Real Decreto 2829/1977 (BOE-A-2025-8109, vigente desde 23/04/2025).",
     {'style': 'List Paragraph'}),
]

# Insertar antes del párrafo "Fin del dossier" usando getprevious()
prev = p_fin._element.getprevious()
if prev is not None and prev.tag == qn('w:p'):
    from docx.text.paragraph import Paragraph
    target_prev = Paragraph(prev, doc.paragraphs[0]._parent)
    insert_chain(target_prev, ins_ref_hhc)
    print('Adición §9 (referencias HHC): OK')
else:
    insert_chain(p_fin, ins_ref_hhc)
    print('Adición §9 (referencias HHC): añadido tras "Fin del dossier"')


# ============================================================
# Guardar
# ============================================================
doc.save(DST)
print(f'\n>>> v4.docx actualizado con INSERCIÓN 14 + renumeración + §9 ampliado')
