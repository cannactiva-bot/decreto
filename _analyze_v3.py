"""Analiza estructura de v3.docx para localizar anclajes de las 12 inserciones."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v3.docx')

# Recorrer body en orden (paragrafos + tablas intercalados)
body = doc.element.body
idx_p = 0
idx_t = 0
order = []
for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        order.append(('P', idx_p))
        idx_p += 1
    elif tag == 'tbl':
        order.append(('T', idx_t))
        idx_t += 1

print(f'TOTAL ELEMENTOS EN BODY: {len(order)} (P:{idx_p} T:{idx_t})\n')

# Anclajes a buscar
anchors = [
    ("INS1", "recurso administrativo en tramitación"),
    ("INS2", "alegando, en sustancia, que las flores"),
    ("INS3", "Equiparación injustificada con el tabaco"),
    ("INS4", "Incoherencia interna del propio ordenamiento sanitario"),
    ("INS5", "Asimetría con los productos sin tabaco que SÍ"),
    ("INS6", "DGT V2221"),
    ("INS6b", "DGT V"),
    ("INS7", "Suiza, que cuenta con el mercado"),
    ("INS8", "C.L.20/2024"),
    ("INS9", "STJUE C-663/18 Kanavape"),
    ("INS10", "de 515 establecimientos especializados en 2023 a 696"),
    ("INS11_marker", "Manifiesto AECANI"),
    ("ERRATA_RECURSO", "Cannabis Hub"),
    ("END_DOC", "Manifiesto AECANI"),
]

for k, (kind, i) in enumerate(order):
    if kind == 'P':
        p = doc.paragraphs[i]
        s = p.style.name if p.style else '-'
        txt = p.text
        for tag, anchor in anchors:
            if anchor.lower() in txt.lower():
                print(f'  >> {tag} found at body[{k}] P[{i}] style={s!r}')
                print(f'     TXT: {txt[:200]}')
                print()

print('\n\n=== HEADINGS ESTRUCTURALES ===')
for k, (kind, i) in enumerate(order):
    if kind == 'P':
        p = doc.paragraphs[i]
        s = p.style.name if p.style else ''
        if s.startswith('Heading'):
            print(f'  body[{k:3d}] P[{i:3d}] {s:12s}: {p.text[:100]}')
