#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para convertir los nuevos documentos (PDF y DOCX) a Markdown
"""

import fitz  # PyMuPDF
import os
import re
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("AVISO: python-docx no instalado. Instalando...")
    import subprocess
    subprocess.run(["pip", "install", "python-docx"], check=True)
    from docx import Document
    HAS_DOCX = True

BASE_DIR = Path(r"C:\Users\micro\Desktop\DECRETO")
MARKDOWN_DIR = BASE_DIR / "markdown"

# Archivos a convertir
FILES_TO_CONVERT = [
    "1- Décret France 30.12.2021.pdf",
    "CBD_as_Food Medicine or Cosmetic_ENG_09102023-Na-22.10.docx",
    "PL Court rulling Kombinat Konopny_Poland (1)-REVIEW.docx",
]

# Buscar el archivo AMCO con wildcard
for f in BASE_DIR.glob("*INFORME AMCO*"):
    FILES_TO_CONVERT.append(f.name)


def pdf_to_markdown(pdf_path, output_path):
    """Convierte PDF a Markdown usando PyMuPDF"""
    doc = fitz.open(pdf_path)
    markdown_lines = []
    
    # Metadatos YAML
    markdown_lines.append("---")
    markdown_lines.append(f"title: {pdf_path.stem}")
    markdown_lines.append(f"source: {pdf_path.name}")
    markdown_lines.append("type: pdf_conversion")
    markdown_lines.append("---\n")
    
    for page_num, page in enumerate(doc, 1):
        markdown_lines.append(f"\n## Pagina {page_num}\n")
        
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if block["type"] == 0:  # Texto
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if text:
                            font_size = span.get("size", 12)
                            is_bold = "bold" in span.get("font", "").lower()
                            
                            # Detectar headings por tamano
                            if font_size >= 16 or (font_size >= 14 and is_bold):
                                text = f"### {text}"
                            elif font_size >= 14:
                                text = f"**{text}**"
                            
                            line_text += text + " "
                    
                    if line_text.strip():
                        markdown_lines.append(line_text.strip())
    
    doc.close()
    
    # Post-procesar
    content = "\n".join(markdown_lines)
    # Limpiar lineas multiples vacias
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return True


def docx_to_markdown(docx_path, output_path):
    """Convierte DOCX a Markdown usando python-docx"""
    doc = Document(docx_path)
    markdown_lines = []
    
    # Metadatos YAML
    markdown_lines.append("---")
    markdown_lines.append(f"title: {docx_path.stem}")
    markdown_lines.append(f"source: {docx_path.name}")
    markdown_lines.append("type: docx_conversion")
    markdown_lines.append("---\n")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        style_name = para.style.name.lower() if para.style else ""
        
        # Detectar headings por estilo
        if "heading 1" in style_name or "titulo 1" in style_name:
            text = f"# {text}"
        elif "heading 2" in style_name or "titulo 2" in style_name:
            text = f"## {text}"
        elif "heading 3" in style_name or "titulo 3" in style_name:
            text = f"### {text}"
        elif "title" in style_name:
            text = f"# {text}"
        
        # Detectar listas
        if text.startswith(("•", "-", "*", "·")):
            text = "- " + text[1:].strip()
        elif re.match(r'^\d+[\.\)]\s', text):
            # Lista numerada
            pass
        
        markdown_lines.append(text)
    
    # Extraer tablas
    for table in doc.tables:
        markdown_lines.append("\n")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            markdown_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                markdown_lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        markdown_lines.append("\n")
    
    content = "\n".join(markdown_lines)
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return True


def main():
    MARKDOWN_DIR.mkdir(exist_ok=True)
    
    print("=" * 60)
    print("CONVERSION DE DOCUMENTOS A MARKDOWN")
    print("=" * 60)
    
    for filename in FILES_TO_CONVERT:
        filepath = BASE_DIR / filename
        
        if not filepath.exists():
            print(f"[X] No encontrado: {filename}")
            continue
        
        # Nombre de salida
        output_name = filepath.stem.replace(" ", "_") + ".md"
        output_path = MARKDOWN_DIR / output_name
        
        print(f"\n[>] Procesando: {filename}")
        
        try:
            if filepath.suffix.lower() == '.pdf':
                pdf_to_markdown(filepath, output_path)
            elif filepath.suffix.lower() == '.docx':
                docx_to_markdown(filepath, output_path)
            else:
                print(f"    Formato no soportado: {filepath.suffix}")
                continue
            
            print(f"    [OK] -> {output_path.name}")
        except Exception as e:
            print(f"    [ERROR] {e}")
    
    print("\n" + "=" * 60)
    print("CONVERSION COMPLETADA")
    print("=" * 60)


if __name__ == "__main__":
    main()
