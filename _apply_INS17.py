"""Aplica INSERCIÓN 17 sobre v4.docx (en local, sin push).

Reemplaza el §4.2.5 actual (version breve de INS15B) por la version expandida
solicitada por el usuario:
- Titulo extendido con subtitulo "...en cuanto a productos de cannabis fuera del ambito agricola".
- Parrafo de apertura scope-limiting.
- Cuerpo (3 parrafos + cita verbatim Evans Medical) ligeramente matizado.
- Parrafo de cierre con sintesis: GMP medicinales + Cosmetica/Novel Food/TPD-GESTABRE-RD 579/2017.

Backup: Dossier_CNMC_AECANI_v4_pre_INS17.docx
"""
import sys, io, shutil, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'
BAK = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4_pre_INS17.docx'
DST = SRC

if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print(f'Backup creado: {os.path.basename(BAK)}')

doc = Document(SRC)

_pStyle_cache = {}

def get_pStyle_xml(style_name):
    if style_name in _pStyle_cache:
        cached = _pStyle_cache[style_name]
        return deepcopy(cached) if cached is not None else None
    pStyle_el = None
    for p in doc.paragraphs:
        try:
            if p.style and p.style.name == style_name:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    found = pPr.find(qn('w:pStyle'))
                    if found is not None:
                        pStyle_el = found
                        break
        except Exception:
            continue
    _pStyle_cache[style_name] = pStyle_el
    return deepcopy(pStyle_el) if pStyle_el is not None else None


def apply_style_xml(p, style_name):
    pStyle_new = get_pStyle_xml(style_name)
    if pStyle_new is None:
        try:
            p.style = style_name
        except KeyError:
            pass
        return
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    existing = pPr.find(qn('w:pStyle'))
    if existing is not None:
        pPr.remove(existing)
    pPr.insert(0, pStyle_new)


def make_run(p, text, italic=False, bold=False):
    r = p.add_run(text)
    if italic:
        r.italic = True
    if bold:
        r.font.bold = True
    return r


def insert_after(target_p, text='', style=None, italic=False, bold=False, indent_quote=False):
    new_p = doc.add_paragraph()
    if style:
        apply_style_xml(new_p, style)
    if text:
        make_run(new_p, text, italic=italic, bold=bold)
    if indent_quote:
        new_p.paragraph_format.left_indent = Inches(0.4)
    target_p._element.addnext(new_p._element)
    return new_p


def insert_chain(target_p, paragraphs):
    cur = target_p
    for spec in paragraphs:
        if isinstance(spec, str):
            text, opts = spec, {}
        else:
            text, opts = spec
        cur = insert_after(cur, text, **opts)
    return cur


# ============================================================
# Localizar el §4.2.5 actual e identificar el rango a borrar
# ============================================================

# Encontrar parrafo con heading "4.2.5"
old_heading = None
for p in doc.paragraphs:
    if p.text.strip().startswith("4.2.5") and "asimetr" in p.text.lower():
        old_heading = p
        break

if old_heading is None:
    print('ERROR: heading 4.2.5 no encontrado'); sys.exit(1)

# Identificar todos los parrafos del bloque §4.2.5 (hasta el siguiente Heading 2 o Heading 3)
to_delete = [old_heading]
sib = old_heading._element.getnext()
while sib is not None:
    tag = sib.tag.split('}')[-1]
    if tag != 'p':  # tabla u otro
        sib = sib.getnext()
        continue
    # Es un <w:p>, comprobar si es Heading
    pPr = sib.find(qn('w:pPr'))
    is_heading = False
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            sval = pStyle.get(qn('w:val'))
            if sval and sval.startswith('Heading'):
                is_heading = True
    if is_heading:
        break
    to_delete.append(sib)
    sib = sib.getnext()

# Convertir a Paragraph wrappers donde sea necesario; ya tenemos elementos lxml.
# Necesitamos guardar el ELEMENTO ANTERIOR al heading para insertar despues.
prev_element = old_heading._element.getprevious()
if prev_element is None:
    print('ERROR: no hay elemento previo al heading'); sys.exit(1)

# Identificar el siguiente heading (el que detuvo el bucle, o None si no se encontro)
next_heading_element = sib  # puede ser None

print(f'Bloque §4.2.5 a borrar: {len(to_delete)} parrafos identificados.')

# Borrar fisicamente
parent = old_heading._element.getparent()
for el in to_delete:
    if isinstance(el, type(old_heading._element)):
        el_to_remove = el
    else:
        el_to_remove = el._element
    parent.remove(el_to_remove)

print(f'Borrados {len(to_delete)} parrafos del §4.2.5 antiguo.')

# Construir el Paragraph wrapper para prev_element
from docx.text.paragraph import Paragraph
target_prev = Paragraph(prev_element, doc.paragraphs[0]._parent)


# ============================================================
# Insertar el §4.2.5 EXPANDIDO solicitado por el usuario
# ============================================================
ins17 = [
    ("4.2.5 La asimetría regulatoria con otros Estados miembros y la doctrina Evans Medical (TJUE C-324/93) en cuanto a productos de cannabis fuera del ámbito agrícola",
     {'style': 'Heading 3'}),

    "Si bien este dossier se centra en los mercados derivados del cáñamo industrial, por similitud y convivencia merece la pena hacer una pequeña mención a aquellos productos de cannabis destinados al mercado medicinal y a cómo se ordenan en España frente a otros Estados miembros.",

    "Los apartados 4 y 5 del artículo 5 del RD 903/2025 imponen al fabricante de preparados estandarizados de cannabis la obtención de licencias específicas conforme al régimen español de fiscalización de estupefacientes (Ley 17/1967) y de sustancias psicótropas. Estas licencias se exigen con independencia del contenido en THC del producto y con independencia de su quimiotipo, y operan adicionalmente respecto de la habilitación farmacéutica general del fabricante.",

    "Para la regulación del cannabis con alto contenido en THC (es decir, fuera del ámbito del cáñamo industrial y de los productos de CBD), otros Estados miembros han optado por regímenes que no exigen tales licencias adicionales. La Medizinal-Cannabisgesetz alemana (MedCanG), a la que se hace referencia adicional en el §4.3 siguiente, sujeta la dispensación de cannabis medicinal a prescripción médica ordinaria sin imponer al fabricante una licencia específica de estupefacientes por el solo hecho de operar con cannabis. Francia opera bajo lógica análoga.",

    "La doctrina del Tribunal de Justicia de la Unión Europea sobre este tipo de asimetría está fijada desde hace tres décadas. La Sentencia de 28 de marzo de 1995, asunto C-324/93 Evans Medical y Macfarlan Smith, dictada precisamente sobre la importación intracomunitaria de un estupefaciente fiscalizado por la Convención Única de 1961, estableció (apartados 31-33):",

    ("«Cuando un Convenio internacional permite que un Estado miembro adopte una medida que es contraria al Derecho comunitario, pero sin obligarle, el Estado miembro debe abstenerse de adoptar dicha medida. Por consiguiente, [...] un Estado miembro debe garantizar la plena eficacia de [la libre circulación de mercancías] dejando inaplicada una práctica nacional contraria, salvo si dicha práctica es necesaria para garantizar la ejecución [...] de obligaciones [...] de un Convenio celebrado con anterioridad.»",
     {'italic': True, 'indent_quote': True}),

    "Aplicada al RD 903/2025: si las licencias de los apartados 4 y 5 del artículo 5 son una medida de fiscalización más estricta de lo estrictamente exigido por las Convenciones de 1961 y 1971 —y la propia existencia de regímenes alternativos en otros Estados miembros que cumplen igualmente con dichas Convenciones es prueba indiciaria de que esa exigencia no es necesaria—, la previsión española constituye medida de efecto equivalente prohibida por el artículo 34 TFUE. Esta dimensión es objeto del recurso 395/2025 ante el Tribunal Supremo (cf. §4.4); su análisis por la CNMC en sede de competencia es complementario y autónomo del proceso jurisdiccional.",

    "En síntesis, un mecanismo de control puede y debe ser proporcional y adecuado, manteniendo suficiente seguridad y garantizando agilidad y flexibilidad en el comercio. En el caso de los productos medicinales, los Estados miembros que han ordenado el sector —Alemania entre ellos— se apoyan en las licencias farmacéuticas y en los protocolos GMP de los laboratorios autorizados, sin necesidad de superponer un régimen adicional de estupefacientes. En el caso de los productos no medicinales y sin riesgo cualificado para la salud pública, los marcos sectoriales europeos vigentes ya controlan adecuadamente la trazabilidad y seguridad de los operadores y de los productos: cosmética (Reglamento (CE) 1223/2009), alimentación (Reglamento (UE) 2015/2283 — Novel Food) y hierbas para fumar (Directiva 2014/40/UE — TPD, plataforma estatal GESTABRE de notificación, Real Decreto 579/2017). El régimen del RD 903/2025 no se inscribe en ninguna de estas lógicas y se configura como una capa adicional de fiscalización que no encuentra equivalente en los Estados miembros de referencia.",
]

cur = target_prev
for spec in ins17:
    if isinstance(spec, str):
        text, opts = spec, {}
    else:
        text, opts = spec
    cur = insert_after(cur, text, **opts)

print(f'INSERCIÓN 17 (§4.2.5 EXPANDIDO): OK -- {len(ins17)} parrafos insertados.')

doc.save(DST)
print(f'\n>>> v4.docx actualizado con §4.2.5 expandido')
