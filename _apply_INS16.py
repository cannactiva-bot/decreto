"""Aplica INSERCIÓN 16 sobre v4.docx (en local, sin push).

- 16A: §7.1 (Interpretación teleológica CN61) - añade Sequeros 2003 como antecedente nacional + contraste con Instrucción Fiscal Antidroga 2021.
- 16B: §4.2.2 (Contradicción con jurisprudencia Kanavape) - añade Francia post-Kanavape: Arrêté 30/12/2021 + CE 444887 (29-dic-2022).
- 16C: §9 - añadir entradas (Sequeros 2003 + CE 444887 + Décret 2022-194 + Arrêté 30/12/2021).

Backup previo: Dossier_CNMC_AECANI_v4_pre_INS16.docx
"""
import sys, io, shutil, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'
BAK = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4_pre_INS16.docx'
DST = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'

if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print(f'Backup creado: {os.path.basename(BAK)}')

doc = Document(SRC)

_pStyle_cache = {}

def get_pStyle_xml(style_name):
    if style_name in _pStyle_cache:
        cached = _pStyle_cache[style_name]
        return deepcopy(cached) if cached is not None else None
    pStyle_el = None
    for p in doc.paragraphs:
        try:
            if p.style and p.style.name == style_name:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    found = pPr.find(qn('w:pStyle'))
                    if found is not None:
                        pStyle_el = found
                        break
        except Exception:
            continue
    _pStyle_cache[style_name] = pStyle_el
    return deepcopy(pStyle_el) if pStyle_el is not None else None


def apply_style_xml(p, style_name):
    pStyle_new = get_pStyle_xml(style_name)
    if pStyle_new is None:
        try:
            p.style = style_name
        except KeyError:
            pass
        return
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    existing = pPr.find(qn('w:pStyle'))
    if existing is not None:
        pPr.remove(existing)
    pPr.insert(0, pStyle_new)


def make_run(p, text, italic=False, bold=False):
    r = p.add_run(text)
    if italic:
        r.italic = True
    if bold:
        r.font.bold = True
    return r


def insert_after(target_p, text='', style=None, italic=False, bold=False, indent_quote=False):
    new_p = doc.add_paragraph()
    if style:
        apply_style_xml(new_p, style)
    if text:
        make_run(new_p, text, italic=italic, bold=bold)
    if indent_quote:
        new_p.paragraph_format.left_indent = Inches(0.4)
    target_p._element.addnext(new_p._element)
    return new_p


def insert_chain(target_p, paragraphs):
    cur = target_p
    for spec in paragraphs:
        if isinstance(spec, str):
            text, opts = spec, {}
        else:
            text, opts = spec
        cur = insert_after(cur, text, **opts)
    return cur


def find_para(text_substring):
    for p in doc.paragraphs:
        if text_substring in p.text:
            return p
    return None


# ============================================================
# INSERCIÓN 16B · §4.2.2 — Francia post-Kanavape (CE 444887)
# Insertamos despues del cierre del §4.2.2 actual.
# ============================================================
p_422_close = find_para("se aparta directamente de esta jurisprudencia")

ins16b = [
    "Otros Estados miembros han ajustado su ordenamiento interno precisamente para dar cumplimiento a la doctrina Kanavape. El caso francés es paradigmático: tras la sentencia del TJUE de 19 de noviembre de 2020, Francia adoptó el Arrêté du 30 décembre 2021 [Anexo 8.4] modificando el régimen del cannabis y autorizando expresamente el cultivo, importación y uso industrial y comercial de variedades de Cannabis sativa L. con contenido inferior al 0,3% de THC. Aquella primera respuesta administrativa francesa, sin embargo, mantenía la prohibición de venta al consumidor de flor y hoja crudas, y fue impugnada por el sector. La cuestión llegó al Conseil d'État —tribunal supremo administrativo francés—, que en su Decisión nº 444887, de 29 de diciembre de 2022 [Anexo 4.6] anuló esa prohibición por desproporcionada, con un razonamiento directamente trasladable al ordenamiento español:",
    ("«Le cannabidiol a des propriétés décontractantes et relaxantes ainsi que des effets anticonvulsivants, [et] il ne présente pas de propriétés psychotropes.»",
     {'italic': True, 'indent_quote': True}),
    ("«La teneur en delta-9-tétrahydrocannabinol des fleurs et de feuilles peut être contrôlée au moyen de tests rapides et peu coûteux.»",
     {'italic': True, 'indent_quote': True}),
    ("«La consommation de variétés de cannabis dont la teneur en THC est inférieure à 0,3 % ne crée pas de risques sanitaires justifiant une interdiction générale et absolue.»",
     {'italic': True, 'indent_quote': True}),
    "El Conseil d'État aplicó así la doctrina Kanavape al supuesto concreto de la flor de cáñamo industrial y concluyó que su prohibición general y absoluta era desproporcionada y carente de justificación sanitaria u orden público. Francia hizo, por tanto, los deberes que el Derecho de la Unión imponía: adaptó su ordenamiento interno (Décret 2022-194 + Arrêté 30/12/2021) y, ante la duda residual sobre la flor cruda, su tribunal supremo administrativo resolvió en sentido alineado con la jurisprudencia del TJUE. España, en cambio, no ha adoptado norma alguna comparable, y el RD 903/2025 mantiene una asimilación general del cannabis al régimen de los estupefacientes que la doctrina del Conseil d'État francés ya ha declarado expresamente injustificada en otro Estado miembro.",
]
if p_422_close:
    insert_chain(p_422_close, ins16b)
    print('INSERCIÓN 16B (§4.2.2 Francia post-Kanavape): OK')
else:
    print('INSERCIÓN 16B: FAIL — anchor "se aparta directamente de esta jurisprudencia" no encontrado')


# ============================================================
# INSERCIÓN 16A · §7.1 — Sequeros 2003 como antecedente nacional
# Insertamos despues del cierre del §7.1, antes del Heading 2 §7.2.
# ============================================================
p_71_close = find_para("incluyendo la propia JIFE")

ins16a = [
    "Esta interpretación teleológica no es novedosa en la doctrina jurídica española. Diecisiete años antes de la sentencia Kanavape, el Fiscal del Tribunal Supremo Fernando Sequeros Sazatornil articuló la misma lógica en su artículo «La venta de semillas de cannabis, de equipos y materiales para su cultivo, así como su propaganda, como actos con trascendencia penal» (Diario La Ley, año XXIV, núm. 5713, 6 de febrero de 2003) [Anexo 5.12], donde defendió que las semillas de cannabis, al carecer de principio activo, quedan excluidas del régimen fiscalizador de la Convención Única de 1961 y, en consecuencia, su comercialización al público no encaja en el tipo penal del artículo 368 del Código Penal:",
    ("«La propaganda de la venta de semillas de cannabis en cantidades para su cultivo doméstico que en principio no excedan de las racionales para su consumo y autoabastecimiento con el mismo fin (incluidas las semillas reproducidas a partir de aquéllas) no podrá considerarse, en modo alguno, constitutiva de delito, al carecer de entidad penal la finalidad perseguida por aquélla.»",
     {'italic': True, 'indent_quote': True}),
    "La doctrina articulada en este artículo —ampliamente invocada por el sector cannábico español durante más de dos décadas y asumida por la Sala Segunda del Tribunal Supremo en una línea jurisprudencial estable (entre otras, STS 1336/2001 de 4 de julio)— constituye antecedente nacional directo del razonamiento que el Tribunal de Justicia de la Unión Europea consagrará en Kanavape: lo que carece de principio activo psicoactivo significativo no es, a los efectos del régimen internacional de fiscalización, estupefaciente. La evolución posterior del Ministerio Fiscal en sentido contrario, particularmente en la Instrucción del Fiscal de Sala Antidroga de junio de 2021 [Anexo 5.13], que endurece el criterio sobre las inflorescencias de cáñamo industrial, se aparta tanto de la propia doctrina interna española como de la jurisprudencia obligatoria Kanavape (2020), generando precisamente la fragmentación documentada en el §5.2(g) del presente dossier.",
]
if p_71_close:
    insert_chain(p_71_close, ins16a)
    print('INSERCIÓN 16A (§7.1 Sequeros 2003): OK')
else:
    print('INSERCIÓN 16A: FAIL — anchor "incluyendo la propia JIFE" no encontrado')


# ============================================================
# INSERCIÓN 16C · §9 — adiciones de referencias
# Anadimos:
#   - Sequeros 2003 al sub-bloque "Doctrina y posiciones administrativas espanolas adicionales"
#   - CE 444887, Decret 2022-194 y Arrete 30/12/2021 al sub-bloque "Jurisprudencia comparada"
# ============================================================

# Anchor 1: el ultimo bullet del sub-bloque "Doctrina y posiciones administrativas..."
# Por ejemplo, despues del bullet sobre Manjon-Cabeza (que esta al final de ese sub-bloque)
p_anchor_doctr = find_para("Cátedra extraordinaria")
if p_anchor_doctr:
    ins16c1 = [
        ("Sequeros Sazatornil, F. — «La venta de semillas de cannabis, de equipos y materiales para su cultivo, así como su propaganda, como actos con trascendencia penal», Diario La Ley, año XXIV, núm. 5713, 6 de febrero de 2003 (Fiscal del Tribunal Supremo).",
         {'style': 'List Paragraph'}),
        ("Instrucción del Fiscal de Sala Antidroga de junio de 2021 sobre flor de cáñamo y CBD.",
         {'style': 'List Paragraph'}),
    ]
    insert_chain(p_anchor_doctr, ins16c1)
    print('INSERCIÓN 16C-1 (§9 doctrina española): OK')

# Anchor 2: ultimo bullet del sub-bloque "Jurisprudencia comparada"
p_anchor_jcomp = find_para("Bundesgericht suizo, Urteil 2C_348/2019")
if p_anchor_jcomp:
    ins16c2 = [
        ("Conseil d'État de Francia, decisión nº 444887, de 29 de diciembre de 2022 (anula la prohibición de venta al consumidor de flor y hoja de cáñamo con < 0,3 % THC).",
         {'style': 'List Paragraph'}),
        ("Décret n° 2022-194, de 17 de febrero de 2022, relativo al cannabis a usos médicos (Francia).",
         {'style': 'List Paragraph'}),
        ("Arrêté du 30 décembre 2021 portant application de l'article R. 5132-86 du code de la santé publique (Francia).",
         {'style': 'List Paragraph'}),
    ]
    insert_chain(p_anchor_jcomp, ins16c2)
    print('INSERCIÓN 16C-2 (§9 jurisprudencia comparada): OK')


doc.save(DST)
print(f'\n>>> v4.docx actualizado con INSERCIÓN 16 (A, B, C)')
