import os
import re
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx', '-q'])
    from docx import Document

def convert_docx_to_markdown(docx_path, output_path):
    """Convert DOCX to Markdown"""
    
    doc = Document(docx_path)
    docx_filename = os.path.basename(docx_path)
    
    markdown_content = f"""---
title: "{docx_filename.replace('.docx', '')}"
source: "{docx_filename}"
converted: "{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
---

"""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_content += "\n"
            continue
        
        style_name = para.style.name.lower() if para.style else ""
        
        # Detect headings
        if 'heading 1' in style_name or 'titulo 1' in style_name:
            markdown_content += f"# {text}\n\n"
        elif 'heading 2' in style_name or 'titulo 2' in style_name:
            markdown_content += f"## {text}\n\n"
        elif 'heading 3' in style_name or 'titulo 3' in style_name:
            markdown_content += f"### {text}\n\n"
        elif 'title' in style_name or 'titulo' in style_name:
            markdown_content += f"# {text}\n\n"
        else:
            # Check if it looks like a heading (short, maybe bold)
            if len(text) < 100 and text.isupper():
                markdown_content += f"## {text}\n\n"
            elif len(text) < 80 and para.runs and any(run.bold for run in para.runs if run.text.strip()):
                # Check if mostly bold
                bold_chars = sum(len(run.text) for run in para.runs if run.bold)
                total_chars = sum(len(run.text) for run in para.runs)
                if total_chars > 0 and bold_chars / total_chars > 0.5:
                    markdown_content += f"### {text}\n\n"
                else:
                    markdown_content += f"{text}\n\n"
            else:
                # Regular paragraph
                markdown_content += f"{text}\n\n"
    
    # Also extract tables
    for table in doc.tables:
        markdown_content += "\n"
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            markdown_content += "| " + " | ".join(cells) + " |\n"
            if i == 0:
                markdown_content += "|" + "|".join(["---"] * len(cells)) + "|\n"
        markdown_content += "\n"
    
    # Clean up
    markdown_content = re.sub(r'\n{4,}', '\n\n\n', markdown_content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    print(f"Conversion completada: {output_path}")
    return True

if __name__ == "__main__":
    conversiones = [
        (r"C:\Users\micro\Downloads\POSITION_PAPER_1pct_THC_ES.docx",
         r"C:\Users\micro\Desktop\DECRETO\position_paper_1pct_thc\Position_Paper_1pct_THC_ES.md"),
        (r"C:\Users\micro\Downloads\ANEXOS_TECNICOS_1pct_THC_ES.docx",
         r"C:\Users\micro\Desktop\DECRETO\position_paper_1pct_thc\Anexos_Tecnicos_1pct_THC_ES.md"),
    ]
    
    for docx_path, output_path in conversiones:
        try:
            if os.path.exists(docx_path):
                convert_docx_to_markdown(docx_path, output_path)
            else:
                print(f"No encontrado: {docx_path}")
        except Exception as e:
            print(f"Error con {docx_path}: {e}")
