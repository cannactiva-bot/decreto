import sys, io, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

def norm(s):
    return ''.join(ch for ch in unicodedata.normalize('NFD', s) if not unicodedata.combining(ch)).lower()

d = Document(r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx')
checks = ['industrial, no medicinal', 'flor de canamo en bruto', 'Clasificacion incorrecta', 'no estupefacientes ni con efectos farmacologicos']

print('=== EN TABLAS ===')
for c in checks:
    n_c = norm(c)
    hits = []
    for ti, tbl in enumerate(d.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                if n_c in norm(cell.text):
                    hits.append(f'tbl[{ti}].row[{ri}].col[{ci}]')
    msg = ', '.join(hits) if hits else 'NO ENCONTRADO'
    print(f'  {c} ::  {msg}')

print()
print('=== EN PARRAFOS ===')
for c in checks:
    n_c = norm(c)
    hits = []
    for pi, p in enumerate(d.paragraphs):
        if n_c in norm(p.text):
            hits.append(f'p[{pi}]')
    msg = ', '.join(hits) if hits else 'NO ENCONTRADO'
    print(f'  {c} ::  {msg}')
