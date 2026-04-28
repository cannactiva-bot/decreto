"""Aplica INSERCIÓN 15 sobre v4.docx (en local, sin push).

Sub-inserciones:
- 15A: §4.3 quinto bullet (asimetría flores/extractos · competencia pura).
- 15B: NUEVO §4.2.5 (Evans Medical doctrine aplicada a las licencias del art. 5).
- 15C: §7.3 tercer bullet (Evans Medical jurisprudencia + cita verbatim).
- 15D: §9 entrada en "Jurisprudencia y régimen internacional".
- 15E: actualización «Ninguna de las dos sentencias» → «Ninguna de las tres sentencias».

Backup previo: Dossier_CNMC_AECANI_v4_pre_INS15.docx
"""
import sys, io, shutil, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'
BAK = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4_pre_INS15.docx'
DST = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'

# Backup
if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print(f'Backup creado: {os.path.basename(BAK)}')

doc = Document(SRC)

_pStyle_cache = {}

def get_pStyle_xml(style_name):
    if style_name in _pStyle_cache:
        cached = _pStyle_cache[style_name]
        return deepcopy(cached) if cached is not None else None
    pStyle_el = None
    for p in doc.paragraphs:
        try:
            if p.style and p.style.name == style_name:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    found = pPr.find(qn('w:pStyle'))
                    if found is not None:
                        pStyle_el = found
                        break
        except Exception:
            continue
    _pStyle_cache[style_name] = pStyle_el
    return deepcopy(pStyle_el) if pStyle_el is not None else None


def apply_style_xml(p, style_name):
    pStyle_new = get_pStyle_xml(style_name)
    if pStyle_new is None:
        try:
            p.style = style_name
        except KeyError:
            pass
        return
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    existing = pPr.find(qn('w:pStyle'))
    if existing is not None:
        pPr.remove(existing)
    pPr.insert(0, pStyle_new)


def make_run(p, text, italic=False, bold=False):
    r = p.add_run(text)
    if italic:
        r.italic = True
    if bold:
        r.font.bold = True
    return r


def insert_after(target_p, text='', style=None, italic=False, bold=False, indent_quote=False):
    new_p = doc.add_paragraph()
    if style:
        apply_style_xml(new_p, style)
    if text:
        make_run(new_p, text, italic=italic, bold=bold)
    if indent_quote:
        new_p.paragraph_format.left_indent = Inches(0.4)
    target_p._element.addnext(new_p._element)
    return new_p


def insert_chain(target_p, paragraphs):
    cur = target_p
    for spec in paragraphs:
        if isinstance(spec, str):
            text, opts = spec, {}
        else:
            text, opts = spec
        cur = insert_after(cur, text, **opts)
    return cur


def find_para(text_substring):
    for p in doc.paragraphs:
        if text_substring in p.text:
            return p
    return None


def replace_text(p, old, new):
    full = p.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    runs = p.runs
    if not runs:
        p.add_run(new_full)
        return True
    for r in runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ''
    return True


# ============================================================
# INSERCIÓN 15B · §4.2 nuevo §4.2.5 (Evans Medical doctrine sobre licencias)
# Insertar tras el ultimo parrafo de §4.2.4
# ============================================================
p_424_close = find_para("La armonización, sea cual sea el nivel, es preferible a la situación actual de vacío normativo")

ins15b = [
    ("4.2.5 La asimetría regulatoria con otros Estados miembros y la doctrina Evans Medical (TJUE C-324/93)",
     {'style': 'Heading 3'}),
    "Los apartados 4 y 5 del artículo 5 del RD 903/2025 imponen al fabricante de preparados estandarizados de cannabis la obtención de licencias específicas conforme al régimen español de fiscalización de estupefacientes (Ley 17/1967) y de sustancias psicótropas. Estas licencias se exigen con independencia del contenido en THC del producto y con independencia de su quimiotipo, y operan adicionalmente respecto de la habilitación farmacéutica general del fabricante.",
    "Otros Estados miembros han optado por regímenes que no exigen tales licencias adicionales. La Medizinal-Cannabisgesetz alemana (MedCanG), a la que se hace referencia adicional en el §4.3 siguiente, sujeta la dispensación de cannabis medicinal a prescripción médica ordinaria sin imponer al fabricante una licencia específica de estupefacientes por el solo hecho de operar con cannabis. Francia opera bajo lógica análoga.",
    "La doctrina del Tribunal de Justicia de la Unión Europea sobre este tipo de asimetría está fijada desde hace tres décadas. La Sentencia de 28 de marzo de 1995, asunto C-324/93 Evans Medical y Macfarlan Smith, dictada precisamente sobre la importación intracomunitaria de un estupefaciente fiscalizado por la Convención Única de 1961, estableció (apartados 31-33):",
    ("«Cuando un Convenio internacional permite que un Estado miembro adopte una medida que es contraria al Derecho comunitario, pero sin obligarle, el Estado miembro debe abstenerse de adoptar dicha medida. Por consiguiente, [...] un Estado miembro debe garantizar la plena eficacia de [la libre circulación de mercancías] dejando inaplicada una práctica nacional contraria, salvo si dicha práctica es necesaria para garantizar la ejecución [...] de obligaciones [...] de un Convenio celebrado con anterioridad.»",
     {'italic': True, 'indent_quote': True}),
    "Aplicada al RD 903/2025: si las licencias de los apartados 4 y 5 del artículo 5 son una medida de fiscalización más estricta de lo estrictamente exigido por las Convenciones de 1961 y 1971 —y la propia existencia de regímenes alternativos en otros Estados miembros que cumplen igualmente con dichas Convenciones es prueba indiciaria de que esa exigencia no es necesaria—, la previsión española constituye medida de efecto equivalente prohibida por el artículo 34 TFUE. Esta dimensión es objeto del recurso 395/2025 ante el Tribunal Supremo (cf. §4.4); su análisis por la CNMC en sede de competencia es complementario y autónomo del proceso jurisdiccional.",
]
if p_424_close:
    insert_chain(p_424_close, ins15b)
    print('INSERCIÓN 15B (§4.2.5 Evans Medical doctrina): OK')
else:
    print('INSERCIÓN 15B: FAIL — anchor "preferible a la situación actual" no encontrado')


# ============================================================
# INSERCIÓN 15A · §4.3 quinto bullet (asimetría flores/extractos · competencia pura)
# Anchor: cierre del bullet 4 ("...en un mercado de demanda creciente.")
# ============================================================
p_43_last = find_para("está renunciando a una ventaja comparativa natural en un mercado de demanda creciente")

ins15a = [
    ("Asimetría material con Alemania en el canal medicinal: exclusión española de la flor. La Medizinal-Cannabisgesetz alemana admite la dispensación medicinal del cannabis en todas sus formas farmacéuticamente válidas, incluida la flor seca (Cannabisblüten), bajo prescripción médica ordinaria a través de la red de oficinas de farmacia comunitarias. El RD 903/2025 español opta por circunscribir su régimen exclusivamente a extractos y preparados estandarizados, dejando la flor fuera del canal medicinal nacional. La consecuencia competitiva es directa: el operador español queda materialmente impedido de suministrar al mercado medicinal interno una de las formas de presentación que el operador alemán comercializa con normalidad. Ambos compiten, sin embargo, en el mismo mercado interior europeo: el operador alemán dispone de un mercado nacional más amplio (extractos + flor) y el español de un mercado nacional reducido (solo extractos), lo que configura una asimetría estructural entre operadores europeos atribuible a la opción regulatoria nacional. Esta asimetría es independiente del juicio de validez sobre la opción restrictiva en sí, y tiene plena dimensión competencial proyectable a la actuación de la CNMC.",
     {'style': 'List Paragraph'}),
]
if p_43_last:
    insert_chain(p_43_last, ins15a)
    print('INSERCIÓN 15A (§4.3 bullet flores/extractos): OK')
else:
    print('INSERCIÓN 15A: FAIL')


# ============================================================
# INSERCIÓN 15C · §7.3 tercer bullet (Evans Medical jurisprudencia)
# Anchor: bullet Kanavape (insertar despues)
# ============================================================
p_73_kana = find_para("STJUE C-663/18 Kanavape (2020)")

ins15c = [
    ("STJUE C-324/93 Evans Medical y Macfarlan Smith, de 28 de marzo de 1995: doctrina fundante sobre los límites del artículo 34 TFUE en presencia de Convenios internacionales sobre estupefacientes. El Tribunal estableció que un Estado miembro debe abstenerse de adoptar medidas de fiscalización más estrictas de las exigidas por las Convenciones de 1961 y 1971 cuando dichas medidas vulneren la libre circulación intracomunitaria. «Cuando un Convenio internacional permite que un Estado miembro adopte una medida que es contraria al Derecho comunitario, pero sin obligarle, el Estado miembro debe abstenerse de adoptar dicha medida» (apartado 32). Doctrina aplicable directamente a las licencias del artículo 5 del RD 903/2025 (cf. §4.2.5).",
     {'style': 'List Paragraph'}),
]
if p_73_kana:
    insert_chain(p_73_kana, ins15c)
    print('INSERCIÓN 15C (§7.3 bullet Evans Medical): OK')
else:
    print('INSERCIÓN 15C: FAIL')


# ============================================================
# INSERCIÓN 15E · actualización «dos» -> «tres» en §7.3
# ============================================================
p_73_close = find_para("Ninguna de las dos sentencias ha sido modulada o revocada")
if p_73_close:
    ok = replace_text(p_73_close, "Ninguna de las dos sentencias", "Ninguna de las tres sentencias")
    print(f'INSERCIÓN 15E (dos→tres): {"OK" if ok else "FAIL"}')


# ============================================================
# INSERCIÓN 15D · §9 entrada en "Jurisprudencia y régimen internacional"
# Anchor: bullet Hammarsten en §9
# ============================================================
p_9_hammarsten = find_para("STJUE C-462/01 (Hammarsten), de 16 de enero de 2003")

ins15d = [
    ("STJUE C-324/93 Evans Medical y Macfarlan Smith, de 28 de marzo de 1995.",
     {'style': 'List Paragraph'}),
]
if p_9_hammarsten:
    insert_chain(p_9_hammarsten, ins15d)
    print('INSERCIÓN 15D (§9 entry Evans Medical): OK')
else:
    print('INSERCIÓN 15D: FAIL')


# ============================================================
# Guardar
# ============================================================
doc.save(DST)
print(f'\n>>> v4.docx actualizado con INSERCIÓN 15 (A,B,C,D,E)')
