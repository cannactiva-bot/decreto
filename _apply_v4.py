"""Aplica las 12 inserciones + erratas de propuestas_v4.md sobre v3.docx -> v4.docx"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v3.docx'
DST = r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx'

doc = Document(SRC)


def make_run(p, text, italic=False, bold=False):
    r = p.add_run(text)
    if italic:
        r.italic = True
    if bold:
        r.font.bold = True
    return r


_pStyle_cache = {}

def get_pStyle_xml(style_name):
    """Devuelve un deepcopy de <w:pStyle> tomado de un parrafo existente con ese estilo."""
    if style_name in _pStyle_cache:
        cached = _pStyle_cache[style_name]
        return deepcopy(cached) if cached is not None else None
    pStyle_el = None
    for p in doc.paragraphs:
        try:
            if p.style and p.style.name == style_name:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    found = pPr.find(qn('w:pStyle'))
                    if found is not None:
                        pStyle_el = found
                        break
        except Exception:
            continue
    _pStyle_cache[style_name] = pStyle_el
    return deepcopy(pStyle_el) if pStyle_el is not None else None


def apply_style_xml(p, style_name):
    """Aplica un estilo via copia XML desde parrafo existente con mismo estilo."""
    pStyle_new = get_pStyle_xml(style_name)
    if pStyle_new is None:
        try:
            p.style = style_name
        except KeyError:
            pass
        return
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    existing = pPr.find(qn('w:pStyle'))
    if existing is not None:
        pPr.remove(existing)
    pPr.insert(0, pStyle_new)


def insert_after(target_p, text='', style=None, italic=False, bold=False, indent_quote=False):
    """Insert a new paragraph immediately after target_p. Returns new paragraph."""
    new_p = doc.add_paragraph()
    if style:
        apply_style_xml(new_p, style)
    if text:
        make_run(new_p, text, italic=italic, bold=bold)
    if indent_quote:
        new_p.paragraph_format.left_indent = Inches(0.4)
    target_p._element.addnext(new_p._element)
    return new_p


def insert_chain(target_p, paragraphs):
    """Insert a sequence of (text, opts) after target_p."""
    cur = target_p
    for spec in paragraphs:
        if isinstance(spec, str):
            text, opts = spec, {}
        else:
            text, opts = spec
        cur = insert_after(cur, text, **opts)
    return cur


def replace_text(p, old, new):
    """Robust paragraph text replacement preserving formatting where possible."""
    full = p.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # Strategy: replace text run-by-run trying to preserve formatting
    runs = p.runs
    if not runs:
        p.add_run(new_full)
        return True
    # If old fits in one run, do surgical replace
    for r in runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # Fallback: clear runs, put all in first
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ''
    return True


# ============================================================
# INSERCIÓN 6 + ERRATA · §6.1 — DGT V2221 -> V2242-22
# ============================================================
p_dgt = doc.paragraphs[166]
old_dgt = "(referencia DGT V2221-CV y otras consultas vinculantes en línea similar)"
new_dgt = ("(referencia DGT V2242-22, de 26 de octubre de 2022, Subdirección General de Impuestos sobre el Consumo, "
           "que aplica a las flores de cáñamo industrial el tipo reducido del 10% del IVA al amparo del artículo "
           "91.Uno.1.8º de la Ley 37/1992)")
ok6 = replace_text(p_dgt, old_dgt, new_dgt)
print(f'INSERCIÓN 6 (DGT): {"OK" if ok6 else "FAIL"}')


# ============================================================
# ERRATA + INSERCIÓN 1 · §4.4 — Recurso 395/2025
# ============================================================
p_rec = doc.paragraphs[131]
old_rec = ("recurso contencioso-administrativo interpuesto por la Asociación Cannabis Hub ante la "
           "Sala Tercera del Tribunal Supremo (recurso administrativo en tramitación, dirección letrada Baño León)")
new_rec = ("recurso contencioso-administrativo nº 395/2025 interpuesto por la Asociación Cannabis Hub "
           "ante la Sala Tercera (Sección Cuarta) del Tribunal Supremo, demanda formalizada el 10 de marzo "
           "de 2026 bajo dirección letrada del despacho Baño León (José María Baño Fos)")
ok1 = replace_text(p_rec, old_rec, new_rec)
print(f'INSERCIÓN 1 + ERRATA recurso §4.4: {"OK" if ok1 else "FAIL"}')


# ============================================================
# ERRATA · §9 — entrada Recurso pendiente TS
# ============================================================
p_rec_ref = doc.paragraphs[259]
old_ref = "Recurso administrativo en tramitación ante la Sala Tercera del Tribunal Supremo (Cannabis Hub vs. Administración del Estado, dirección letrada Baño León)."
new_ref = "Recurso contencioso-administrativo nº 395/2025, Sala Tercera-Sección Cuarta del Tribunal Supremo, Cannabis Hub vs. Administración del Estado; demanda formalizada el 10 de marzo de 2026 (dirección letrada Baño León)."
ok_ref = replace_text(p_rec_ref, old_ref, new_ref)
print(f'ERRATA §9 recurso: {"OK" if ok_ref else "FAIL"}')


# ============================================================
# INSERCIÓN 10 · §3.4 — colapso del cultivo
# ============================================================
p_88 = doc.paragraphs[88]  # último bullet "Crecimiento sostenido..."
ins10_paras = [
    ("Colapso paralelo del cultivo nacional: la superficie declarada de cáñamo industrial en España cayó "
     "de 688 hectáreas en 2020 a 62 hectáreas en 2024 (−91%), según datos FEGA y Cannamonitor "
     "(edición octubre 2025). En 2024 se produjeron 51,3 toneladas de materia prima medicinal en territorio "
     "español, íntegramente exportadas. Es decir: los laboratorios farmacéuticos españoles autorizados por "
     "el RD 903/2025 se ven obligados a importar materia prima desde terceros países o desde otros Estados "
     "miembros mientras la producción nacional se exporta en su totalidad.",
     {'style': 'List Paragraph'}),
    ("LA SEGUNDA CONSTATACIÓN. Coexisten dos colapsos simétricos: el establecimiento de venta minorista "
     "crece (+35% en dos años) mientras la producción agrícola nacional se desploma (−91% en cuatro años). "
     "Ambos fenómenos comparten una misma causa: la inexistencia de un marco regulatorio coherente que "
     "ordene el sector y permita conectar la oferta agrícola nacional con la demanda nacional —medicinal "
     "y no medicinal— ya existente.",
     {'bold': True}),
]
insert_chain(p_88, ins10_paras)
print('INSERCIÓN 10 (cultivo -91%): OK')


# ============================================================
# INSERCIÓN 2 · §5.1 — PPT Verdejo + Doc. Consenso + Lista positivos
# ============================================================
p_140 = doc.paragraphs[140]  # "...sustancias estupefacientes —lo cual contradice todo lo expuesto..."
ins2_paras = [
    "Esta posición administrativa ha sido sostenida públicamente por la propia Jefa de Área de Prevención y Control del Tabaquismo del Ministerio de Sanidad. En la presentación oficial «Impacto del tabaco y las nuevas formas de consumo en la infancia y la adolescencia», expuesta en la I Jornada de Escuelas Promotoras de Salud (15 de octubre de 2024, Auditorio Ernest Lluch, Madrid), publicada en sanidad.gob.es, se afirma literalmente:",
    ("«Productos a base de hierbas para fumar […] No se autoriza dentro de este grupo de productos: cáñamo, cannabis, cbd, o cualquier otra planta con efectos psicoactivos.» (diapositiva 15) [Anexo 9.1]",
     {'italic': True, 'indent_quote': True}),
    "Y en la diapositiva 21:",
    ("«Productos con cannabis y derivados con CBD y bajo contenido en THC (<0,3%): No están autorizados por el Ministerio de Sanidad ni por la Agencia del Medicamento. Publicidad falsa.» [Anexo 9.1]",
     {'italic': True, 'indent_quote': True}),
    "Esta exclusión administrativa contrasta frontalmente con el Documento de Consenso sobre la aplicación de la Ley 28/2005 (Grupo de responsables de tabaquismo MSCBS-CCAA, revisado por la Comisión de Salud Pública el 13 de noviembre de 2019), que en su sección 13 establece que «la venta de productos a base de hierbas para fumar es libre salvo para el caso de la prohibición de venta a menores» [Anexo 9.2]. La lista positiva oficial de productos a base de hierbas para fumar publicada por el propio Ministerio (Lista_positivos_Hierbas.pdf) no incluye, a fecha de cierre del presente dossier, ningún producto derivado del cáñamo industrial [Anexo 9.3].",
    "El conjunto evidencia una exclusión discrecional sin base normativa expresa, articulada al margen del régimen general de venta libre, y contradictoria con el artículo 2.4 de la Directiva 2014/40/UE TPD (cuya definición de hierbas para fumar se basa en la ausencia de tabaco, no en la ausencia de cáñamo) y con la jurisprudencia Kanavape (TJUE C-663/18).",
]
insert_chain(p_140, ins2_paras)
print('INSERCIÓN 2 (PPT Verdejo): OK')


# ============================================================
# INSERCIÓN 3 · §5.2(a) — Sentenza 83/2015 Italia
# ============================================================
# Subseccion (a) Heading 3 en P[148]; el parrafo de contenido es P[149]; luego Heading 3 (b) en P[150]
# Insertar despues de P[149]
p_149 = doc.paragraphs[149]
ins3_paras = [
    "El precedente constitucional italiano —país con la red más densa de monopolio del tabaco en Europa— resulta directamente trasladable. La Corte Costituzionale, en su Sentenza n. 83 de 14 de abril de 2015 (depositada el 15 de mayo de 2015, ECLI:IT:COST:2015:83), declaró la ilegitimidad constitucional de la extensión del régimen del tabaco a los líquidos para cigarrillo electrónico sin nicotina, con el siguiente razonamiento (Considerato in diritto, punto 5.1):",
    ("«Appare quindi del tutto irragionevole l'estensione, operata dalla disposizione censurata, del regime amministrativo e tributario proprio dei tabacchi anche al commercio di liquidi aromatizzati e di dispositivi per il relativo consumo, i quali non possono essere considerati succedanei del tabacco.» [Anexo 9.8]",
     {'italic': True, 'indent_quote': True}),
    "Y al cualificar la justificación específica del régimen del tabaco que falta en otros productos:",
    ("«il regime fiscale dell'accisa con riferimento al mercato dei tabacchi, trova la sua giustificazione nel disfavore nei confronti di un bene riconosciuto come gravemente nocivo per la salute (...), tale presupposto non è ravvisabile in relazione al commercio di prodotti contenenti "
     "‘altre sostanze’, diverse dalla nicotina.» [Anexo 9.8]",
     {'italic': True, 'indent_quote': True}),
    "La Sentenza n. 240 de 15 de noviembre de 2017 (ECLI:IT:COST:2017:240) confirmó la doctrina a contrario, validando un régimen diferenciado y proporcional [Anexo 9.9]. Por argumento a fortiori: si la equiparación al tabaco es inconstitucional para un líquido de vapeo sin nicotina, lo es con mayor razón para una flor agrícola no manufacturada, sin tabaco y sin nicotina.",
]
insert_chain(p_149, ins3_paras)
print('INSERCIÓN 3 (Sentenza 83/2015): OK')


# ============================================================
# INSERCIÓN 4 · §5.2(d) — Circular 3/2020 Comisionado
# ============================================================
# Tras Insercion 3 los indices se desplazaron. Reabrir el doc no es necesario porque
# python-docx mantiene la referencia, pero el indice numerico cambia.
# Localizamos por TEXTO en lugar de por indice.
def find_para(text_substring, exact_start=False):
    for p in doc.paragraphs:
        if exact_start:
            if p.text.startswith(text_substring):
                return p
        else:
            if text_substring in p.text:
                return p
    return None

# (d) Heading 3 -> contenido -> (e) Heading 3
p_d_content = find_para("El propio Real Decreto 579/2017 prohíbe en su articulado la utilización de CBD")
ins4_paras = [
    "A esta incoherencia interna se suma la posición pública del propio Comisionado para el Mercado de Tabacos, expresada en la Circular 3/2020, de 27 de noviembre [Anexo 9.4], que estableció:",
    ("«Por todo lo anterior, el Comisionado para el Mercado de Tabacos desea advertir a todos los operadores de que la comercialización de productos de cannabis, independientemente de su contenido en tetrahidrocannabinol (THC), se encuentra prohibida por la Ley.»",
     {'italic': True, 'indent_quote': True}),
    "La paradoja regulatoria es completa: el órgano gestor del monopolio de las expendedurías declara expresamente prohibida la comercialización de productos de cannabis en su red, mientras simultáneamente se propone, vía TRIS 2025/0044/ES, canalizar precisamente esos mismos productos a través de esa misma red. La posición de la Administración española sobre este punto es, en su literalidad, internamente contradictoria entre dos órganos del mismo Ministerio (Hacienda).",
]
if p_d_content:
    insert_chain(p_d_content, ins4_paras)
    print('INSERCIÓN 4 (Circular 3/2020): OK')
else:
    print('INSERCIÓN 4: FAIL (anchor not found)')


# ============================================================
# INSERCIÓN 5 · §5.2 nueva letra (g) — fragmentación STS
# ============================================================
# Insertar entre el parrafo de (f) y el Heading 2 "5.3 La pregunta-respuesta"
p_f_content = find_para("Los cigarrillos electrónicos sin nicotina, los productos a base de hierbas calentadas")
ins5_paras = [
    ("(g) Fragmentación jurisprudencial dentro del propio Tribunal Supremo y los tribunales inferiores",
     {'style': 'Heading 3'}),
    "La conducta consistente en la comercialización de flor de cáñamo industrial con contenido THC inferior al 0,3% recibe tratamientos jurisdiccionales radicalmente divergentes en función del juzgado que conoce del asunto. La línea condenatoria del Tribunal Supremo (entre otras: STS 288/2023, STS 177/2024, STS 583/2024, STS 678/2024) coexiste con una línea absolutoria nutrida en juzgados penales y audiencias provinciales (Juzgado de lo Penal nº 8 de Valencia, 2021 —primera resolución española que invoca expresamente Kanavape—; AP Teruel 53/2024 de 5 de julio; TSJ Madrid, marzo 2024; archivos en País Vasco y Las Palmas). La misma conducta —comercializar flor de cáñamo con THC <0,3%— puede traducirse, según el órgano que conozca, en archivo o en condena de hasta cuatro años de prisión. Esta fragmentación judicial es por sí misma indicio de defecto regulatorio aplicable al artículo 28 LGUM.",
]
if p_f_content:
    insert_chain(p_f_content, ins5_paras)
    print('INSERCIÓN 5 (fragmentación STS): OK')
else:
    print('INSERCIÓN 5: FAIL (anchor not found)')


# ============================================================
# INSERCIÓN 7 · §6.2(b) — BGer 2C_348/2019 Suiza
# ============================================================
p_suiza = find_para("Suiza, que cuenta con el mercado de CBD-cannabis light")
ins7_paras = [
    "Este criterio fue judicialmente confirmado por el Tribunal Federal Suizo (Bundesgericht) en su sentencia 2C_348/2019, de 29 de enero de 2020 [Anexo 9.10] (con las gemelas 2C_350/2019 y 2C_402/2019), que revocó la decisión administrativa de la Oberzolldirektion (Merkblatt OZD de 22 de febrero de 2017) que había sometido las flores de cáñamo CBD al impuesto sobre el tabaco. El holding (E. 4.8) es taxativo:",
    ("«Demzufolge fehlt es im TStG und der TStV an einer gesetzlichen Grundlage, um Cannabisblüten der Tabaksteuer zu unterwerfen.» (En consecuencia, falta en la Ley del impuesto sobre el tabaco y en su reglamento de desarrollo una base legal para someter las flores de cannabis al impuesto sobre el tabaco.)",
     {'italic': True, 'indent_quote': True}),
    "Y en la ratio (E. 4.7):",
    ("«Aus Sicht des Konsumenten sind Cannabisblüten gerade kein Ersatz für herkömmliche Tabakfabrikate, sondern befriedigen andere Bedürfnisse. […] Bei Cannabisblüten handelt es sich, auch wenn sie unter anderem geraucht werden, um ein Produkt mit speziellen Eigenschaften, welches zu Tabakprodukten nicht in einem Substitutionsverhältnis steht.» (Desde el punto de vista del consumidor, las flores de cannabis no son sustituto de los productos clásicos del tabaco, sino que satisfacen otras necesidades. […] Las flores de cannabis, aun cuando entre otras cosas se fumen, son un producto con propiedades especiales que no se halla en relación de sustitución con los productos del tabaco.)",
     {'italic': True, 'indent_quote': True}),
    "Como consecuencia directa de la sentencia, la Oficina Federal de Aduanas y Seguridad de las Fronteras (BAZG) procedió a la devolución de aproximadamente 33 millones de francos suizos a los productores afectados.",
]
if p_suiza:
    insert_chain(p_suiza, ins7_paras)
    print('INSERCIÓN 7 (BGer Suiza): OK')
else:
    print('INSERCIÓN 7: FAIL')


# ============================================================
# INSERCIÓN 8 · §7.2 — WHO ECDD CBD 2018
# ============================================================
p_jife = find_para("La autoridad de Naciones Unidas competente para interpretar y vigilar")
# Insertar al final de §7.2 (antes del Heading 2 §7.3). El parrafo actual tras p_jife es uno solo:
# "Esta interpretación es coherente con la Decisión (UE) 2021/3..."
p_jife_close = find_para("Esta interpretación es coherente con la Decisión (UE) 2021/3")
ins8_paras = [
    "La posición de la JIFE se inserta en una línea de evolución del régimen internacional de control que parte del «Cannabidiol (CBD) — Critical Review Report» del Comité de Expertos en Farmacodependencia de la Organización Mundial de la Salud (ECDD), 40ª reunión, Ginebra, junio de 2018 [Anexo 9.7]. Este informe estableció con base en la evidencia científica disponible que el cannabidiol «no muestra potencial de abuso ni de dependencia» y que «no se asocia a ninguna problemática de salud pública», fundamentando así la subsiguiente Decisión (UE) 2021/3 del Consejo y la posición progresiva de la JIFE.",
]
if p_jife_close:
    insert_chain(p_jife_close, ins8_paras)
    print('INSERCIÓN 8 (WHO ECDD 2018): OK')
else:
    print('INSERCIÓN 8: FAIL')


# ============================================================
# INSERCIÓN 9 · §7.3 — Conclusiones AG Tanchev
# ============================================================
# Insertar al final de §7.3 (despues del parrafo "Ninguna de las dos sentencias ha sido modulada o revocada...")
p_tjue_close = find_para("Ninguna de las dos sentencias ha sido modulada o revocada")
ins9_paras = [
    "Las Conclusiones del Abogado General Sr. Evgeni Tanchev, presentadas el 14 de mayo de 2020 en el asunto C-663/18 [Anexo 9.6], son particularmente explícitas en cuanto al estándar probatorio que debe cumplir un Estado miembro para justificar restricciones nacionales sobre el CBD bajo el artículo 36 TFUE:",
    ("«Habida cuenta de la información facilitada al Tribunal de Justicia, difícilmente cabe considerar que el Gobierno francés haya identificado de forma clara los efectos nocivos, concretamente psicotrópicos, que entraña el uso de aceite de CBD en cigarrillos electrónicos, y aún menos que haya efectuado un análisis exhaustivo del riesgo para la salud basado en los datos científicos más fiables de que se disponga y en los resultados más recientes de la investigación internacional.» (Conclusiones AG Tanchev, §83)",
     {'italic': True, 'indent_quote': True}),
    "El análisis del Abogado General es íntegramente proyectable al caso español: la Administración española no ha aportado, en sede de motivación de las decisiones administrativas que bloquean al sector, un análisis del riesgo basado en datos científicos comparable al estándar exigido por el Tribunal de Justicia.",
]
if p_tjue_close:
    insert_chain(p_tjue_close, ins9_paras)
    print('INSERCIÓN 9 (AG Tanchev): OK')
else:
    print('INSERCIÓN 9: FAIL')


# ============================================================
# INSERCIÓN 13 · §5.1 — Régimen aduanero Luxemburgo (Douanes LU)
# ============================================================
# Tras la cita del IPN/CNMC/040/24 al final de §5.1, antes del Heading 2 §5.2
p_ipn = find_para("La medida establecería de facto una prohibición a la comercialización")
ins13_paras = [
    "A modo ilustrativo de la práctica administrativa de un Estado miembro vecino, la Administration des Douanes et Accises de Luxemburgo ha publicado en su sitio oficial (douanes.public.lu), bajo la rúbrica «Tabacs manufacturés · Produits à base de cannabis», un régimen específico vigente desde el 19 de marzo de 2019 para los productos a base de cannabis con contenido en THC inferior al 1% destinados a ser fumados o vaporizados, comprendiendo expresamente:",
    ("«les fleurs de chanvre ou les parties de plantes susceptibles d'être fumées ou vaporisées (THC < 1%) ; la résine de chanvre (THC < 1%) ; les cigarettes composées exclusivement de chanvre (THC < 1%).» [Anexo 8.5]",
     {'italic': True, 'indent_quote': True}),
    "Es decir: en un Estado miembro de la Unión Europea la propia administración aduanera identifica las flores y demás partes de la planta de cáñamo industrial con contenido inferior al 1% de THC como categoría comercial autónoma y diferenciada, con un régimen administrativo específico, mientras que en España la administración sanitaria no admite siquiera su notificación bajo el régimen general de hierbas para fumar. La asimetría con la situación española es, por sí misma, expresiva, y se refuerza al constatar que el umbral aplicado por Luxemburgo (1% THC) coincide con la definición de quimiotipo CBD de la Farmacopea Europea (Ph. Eur. 3028) y con la legislación checa, en línea con la práctica seguida por aproximadamente cuarenta jurisdicciones a escala mundial.",
]
if p_ipn:
    insert_chain(p_ipn, ins13_paras)
    print('INSERCIÓN 13 (Douanes LU): OK')
else:
    print('INSERCIÓN 13: FAIL (anchor not found)')


# ============================================================
# INSERCIÓN 11 · §9 — adiciones documentación de referencia
# ============================================================
# Insertar al final de la ultima sub-seccion de §9 (Documentacion AECANI complementaria)
# Necesitamos insertar 3 nuevos sub-bloques: Heading 2 + bullets, antes del cierre del documento.
# El cierre suele ser "— Fin del dossier —" o similar. Buscamos.
p_fin = find_para("Fin del dossier")
if p_fin is None:
    # Si no hay marca de fin, insertamos despues del ultimo bullet AECANI
    # Buscamos la ultima entrada de "Documentacion AECANI complementaria"
    p_fin = find_para("Paquete de preguntas parlamentarias al Gobierno")
    if p_fin is None:
        p_fin = find_para("Manifiesto AECANI")

ins11_blocks = [
    ("Doctrina y posiciones administrativas españolas adicionales", {'style': 'Heading 2'}),
    ("Consulta vinculante de la Dirección General de Tributos V2242-22, de 26 de octubre de 2022 (IVA aplicable a las flores de cáñamo industrial).", {'style': 'List Paragraph'}),
    ("Circular 3/2020, de 27 de noviembre, del Comisionado para el Mercado de Tabacos sobre comercialización de productos derivados del cannabis en expendedurías.", {'style': 'List Paragraph'}),
    ("Documento de Consenso sobre aplicación de la Ley 28/2005 (Grupo de responsables de tabaquismo MSCBS-CCAA, revisado por Comisión de Salud Pública 13/11/2019).", {'style': 'List Paragraph'}),
    ("Lista positiva oficial de productos a base de hierbas para fumar (Ministerio de Sanidad).", {'style': 'List Paragraph'}),
    ("Presentación oficial «Impacto del tabaco y las nuevas formas de consumo en la infancia y la adolescencia», Susana Verdejo Fernández (Jefa de Área de Prevención y Control del Tabaquismo, Ministerio de Sanidad), I Jornada de Escuelas Promotoras de Salud, 15/10/2024.", {'style': 'List Paragraph'}),
    ("Informe jurídico de Araceli Manjón-Cabeza Olmeda (Catedrática de Derecho Penal · directora de la Cátedra extraordinaria «Drogas Siglo XXI» — Universidad Complutense de Madrid), encargado por la Asociación Cannabis Hub, sobre la regulación internacional del cáñamo industrial y del CBD (firma digital 14/07/2024).", {'style': 'List Paragraph'}),

    ("Jurisprudencia comparada en materia de cáñamo y régimen del tabaco", {'style': 'Heading 2'}),
    ("Conclusiones del Abogado General Sr. E. Tanchev en el asunto TJUE C-663/18 (14/05/2020).", {'style': 'List Paragraph'}),
    ("Corte Costituzionale italiana, Sentenza n. 83 de 14 de abril de 2015 (ECLI:IT:COST:2015:83).", {'style': 'List Paragraph'}),
    ("Corte Costituzionale italiana, Sentenza n. 240 de 15 de noviembre de 2017 (ECLI:IT:COST:2017:240).", {'style': 'List Paragraph'}),
    ("Bundesgericht suizo, Urteil 2C_348/2019, 2C_350/2019 y 2C_402/2019, de 29 de enero de 2020.", {'style': 'List Paragraph'}),

    ("Apoyo parlamentario", {'style': 'Heading 2'}),
    ("Proposición no de Ley en Comisión de Agricultura, presentada por los Grupos Parlamentarios Republicano, Sumar y EH Bildu, instando a la AEMPS a declarar la conformidad de su práctica administrativa con la jurisprudencia Kanavape (TJUE C-663/18).", {'style': 'List Paragraph'}),
]
if p_fin:
    # insertamos ANTES de p_fin: para eso, insertamos cada bloque despues del anterior
    # truco: insertamos el primero ANTES de p_fin, luego en cadena
    # Como nuestro helper insert_after solo añade despues de target,
    # insertamos despues del PARRAFO PREVIO al de fin (su anterior).
    prev = p_fin._element.getprevious()
    if prev is not None:
        # Crear un parrafo dummy para usar como target
        # Mas facil: buscar el ultimo bullet AECANI antes de p_fin y arrancar desde alli
        # En lugar de complicarlo, insertamos despues del parrafo previo accediendo via el body
        # Solucion: encadenamos las inserciones colgando del previous element
        # Convertimos prev a paragraph object si es un <w:p>
        from docx.text.paragraph import Paragraph
        if prev.tag == qn('w:p'):
            target_prev = Paragraph(prev, doc.paragraphs[0]._parent)
            insert_chain(target_prev, ins11_blocks)
            print('INSERCIÓN 11 (§9 adiciones): OK (insertado antes de "Fin del dossier")')
        else:
            insert_chain(p_fin, ins11_blocks)  # fallback: despues
            print('INSERCIÓN 11 (§9 adiciones): OK (despues del marcador de fin)')
    else:
        insert_chain(p_fin, ins11_blocks)
        print('INSERCIÓN 11 (§9 adiciones): OK (insertado despues)')
else:
    # ultimo recurso: añadir al final del documento
    cur = doc.paragraphs[-1]
    insert_chain(cur, ins11_blocks)
    print('INSERCIÓN 11 (§9 adiciones): añadido al final del documento')


# ============================================================
# Guardar
# ============================================================
doc.save(DST)
print(f'\n>>> Guardado: {DST}')
