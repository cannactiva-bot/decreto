# -*- coding: utf-8 -*-
"""
Convierte INFORME_RECURSO_RD_CANNABIS_v3.md a DOCX
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_styles(doc):
    """Crea estilos personalizados para el documento"""
    styles = doc.styles
    
    # Estilo para titulo principal
    if 'TituloPrincipal' not in [s.name for s in styles]:
        style = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(18)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilo para citas
    if 'Cita' not in [s.name for s in styles]:
        style = styles.add_style('Cita', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(10)
        style.font.italic = True
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.right_indent = Inches(0.5)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

def parse_table(lines, start_idx):
    """Parsea una tabla markdown y devuelve los datos"""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    
    if len(table_lines) < 2:
        return None, start_idx
    
    # Parsear filas
    rows = []
    for line in table_lines:
        if '---' in line:  # Skip separator
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    return rows, i

def add_table_to_doc(doc, rows):
    """Agrega una tabla al documento"""
    if not rows or len(rows) < 1:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                # Limpiar markdown del texto
                clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
                cell.text = clean_text
                
                # Primera fila en negrita (encabezado)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()

def convert_md_to_docx(md_path, docx_path):
    """Convierte un archivo markdown a docx"""
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    create_styles(doc)
    
    # Configurar margenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # Bloques de codigo
        if line.strip().startswith('```'):
            if in_code_block:
                # Fin del bloque
                if code_content:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    for code_line in code_content:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Tablas
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            rows, new_i = parse_table(lines, i)
            if rows:
                add_table_to_doc(doc, rows)
            i = new_i
            continue
        
        # Encabezados
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Remove links
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        if line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Linea horizontal
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Citas (blockquotes)
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            # Recoger lineas adicionales de la cita
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                i += 1
                quote_text += ' ' + lines[i].strip()[1:].strip()
            
            # Limpiar markdown
            quote_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', quote_text)
            quote_text = re.sub(r'\*([^*]+)\*', r'\1', quote_text)
            
            p = doc.add_paragraph(quote_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.3)
            p.italic = True
            for run in p.runs:
                run.italic = True
            i += 1
            continue
        
        # Linea vacia
        if not line.strip():
            i += 1
            continue
        
        # Parrafo normal
        text = line.strip()
        
        # Limpiar enlaces markdown
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Procesar negrita e italica
        p = doc.add_paragraph()
        
        # Patron para negrita e italica
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)
        
        i += 1
    
    doc.save(docx_path)
    print(f"Documento guardado en: {docx_path}")

if __name__ == '__main__':
    md_path = r'C:\Users\micro\Desktop\DECRETO\INFORME_RECURSO_RD_CANNABIS_v3.md'
    docx_path = r'C:\Users\micro\Desktop\DECRETO\INFORME_RECURSO_RD_CANNABIS_v3.docx'
    convert_md_to_docx(md_path, docx_path)
