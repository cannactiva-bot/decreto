import sys, io, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

def norm(s):
    return ''.join(ch for ch in unicodedata.normalize('NFD', s) if not unicodedata.combining(ch)).lower()

d = Document(r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx')
print(f'PARRAFOS: {len(d.paragraphs)} | TABLAS: {len(d.tables)} | CARS: {sum(len(p.text) for p in d.paragraphs)}')

checks = [
    ('Sequeros §7.1', 'Sequeros Sazatornil'),
    ('Sequeros cita', 'no podra considerarse, en modo alguno'),
    ('Francia §4.2.2 Conseil d\'Etat', 'Conseil d'+chr(0x2019)+'Etat'),
    ('Francia 444887', '444887'),
    ('Cannabidiol psychotropes', 'ne presente pas de proprietes'),
    ('§4.2.5 expandido scope', 'Si bien este dossier se centra'),
    ('§4.2.5 GESTABRE', 'GESTABRE'),
    ('§4.3 Cannabisbluten', 'Cannabisbluten'),
    ('§7.3 tres sentencias', 'tres sentencias'),
    ('Evans Medical bullet', 'Evans Medical'),
    ('§9 Sequeros entry', 'Diario La Ley'),
    ('§9 CE 444887 entry', 'Conseil d'+chr(0x2019)+'Etat de Francia'),
]
for label, c in checks:
    n_c = norm(c)
    cnt = sum(1 for p in d.paragraphs if n_c in norm(p.text))
    status = 'OK' if cnt > 0 else 'MISSING'
    print(f'  [{status}] {label}: {cnt}x')
