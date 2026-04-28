import sys, io, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

def norm(s):
    return ''.join(ch for ch in unicodedata.normalize('NFD', s) if not unicodedata.combining(ch)).lower()

d = Document(r'C:\Users\micro\Desktop\cnmc\Dossier_CNMC_AECANI_v4.docx')
print(f'PARRAFOS: {len(d.paragraphs)} | TABLAS: {len(d.tables)}')
print(f'CARS: {sum(len(p.text) for p in d.paragraphs)}')

checks = [
    'Evans Medical',
    'C-324/93',
    '4.2.5 La asimetria',
    'Asimetria material con Alemania',
    'Cannabisbluten',
    'tres sentencias',
    'apartados 31-33',
]
for c in checks:
    n_c = norm(c)
    cnt = sum(1 for p in d.paragraphs if n_c in norm(p.text))
    print(f'  {c} :: {cnt}x')
