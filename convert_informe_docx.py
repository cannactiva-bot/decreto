#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convierte el informe markdown a Word docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level):
    """Agrega un heading con el nivel apropiado"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_with_formatting(doc, text):
    """Agrega un parrafo con formato basico"""
    # Procesar negritas
    p = doc.add_paragraph()
    
    # Patron para negritas **texto**
    pattern = r'\*\*([^*]+)\*\*'
    parts = re.split(pattern, text)
    
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:  # Es texto en negrita
            run.bold = True
    
    return p

def add_table_from_markdown(doc, lines):
    """Convierte una tabla markdown a tabla Word"""
    # Filtrar lineas vacias y separadores
    table_lines = [l for l in lines if l.strip() and not l.strip().startswith('|--') and not re.match(r'^\|[-:\s|]+\|$', l.strip())]
    
    if not table_lines:
        return
    
    # Obtener numero de columnas
    first_line = table_lines[0]
    cols = [c.strip() for c in first_line.split('|')[1:-1]]
    num_cols = len(cols)
    
    if num_cols == 0:
        return
    
    # Crear tabla
    table = doc.add_table(rows=len(table_lines), cols=num_cols)
    table.style = 'Table Grid'
    
    for row_idx, line in enumerate(table_lines):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        for col_idx, cell_text in enumerate(cells):
            if col_idx < num_cols:
                # Limpiar formato markdown
                cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                table.rows[row_idx].cells[col_idx].text = cell_text
    
    # Hacer header en negrita
    if table.rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

def convert_md_to_docx(md_path, docx_path):
    """Convierte archivo markdown a docx"""
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Saltar lineas vacias
        if not stripped:
            i += 1
            continue
        
        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            add_heading(doc, text, 1)
            i += 1
            continue
        
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            add_heading(doc, text, 2)
            i += 1
            continue
        
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            add_heading(doc, text, 3)
            i += 1
            continue
        
        # Lineas horizontales
        if stripped == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Tablas
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, table_lines)
            doc.add_paragraph()  # Espacio despues de tabla
            continue
        
        # Citas (blockquotes)
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.style = 'Quote' if 'Quote' in doc.styles else 'Normal'
            
            # Procesar negritas en cita
            pattern = r'\*\*([^*]+)\*\*'
            parts = re.split(pattern, quote_text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.italic = True
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue
        
        # Listas
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Procesar negritas
            pattern = r'\*\*([^*]+)\*\*'
            parts = re.split(pattern, text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue
        
        # Listas numeradas
        if re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            p = doc.add_paragraph(style='List Number')
            pattern = r'\*\*([^*]+)\*\*'
            parts = re.split(pattern, text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue
        
        # Bloques de codigo
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Saltar el cierre ```
            
            # Agregar como parrafo con fuente monoespaciada
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.3)
            continue
        
        # Parrafo normal
        add_paragraph_with_formatting(doc, stripped)
        i += 1
    
    # Guardar documento
    doc.save(docx_path)
    print(f"Documento guardado: {docx_path}")

if __name__ == '__main__':
    md_path = r'C:\Users\micro\Desktop\DECRETO\INFORME_RECURSO_RD_CANNABIS_ESTRUCTURA_DEFINITIVA.md'
    docx_path = r'C:\Users\micro\Desktop\DECRETO\INFORME_RECURSO_RD_CANNABIS_ESTRUCTURA_DEFINITIVA.docx'
    convert_md_to_docx(md_path, docx_path)
