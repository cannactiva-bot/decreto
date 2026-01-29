#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convierte INFORME_RECURSO_RD_CANNABIS.md a formato Word (.docx)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_document_styles(doc):
    """Configura los estilos del documento"""
    # Estilo Normal
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Titulos
    for i in range(1, 5):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = 'Calibri'
            style.font.color.rgb = RGBColor(0, 51, 102)
            if i == 1:
                style.font.size = Pt(18)
                style.font.bold = True
            elif i == 2:
                style.font.size = Pt(14)
                style.font.bold = True
            elif i == 3:
                style.font.size = Pt(12)
                style.font.bold = True
            else:
                style.font.size = Pt(11)
                style.font.bold = True

def add_code_block(doc, text):
    """Agrega un bloque de codigo con formato"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Agregar sombreado
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # Agregar borde
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    
    # Sombreado
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

def add_quote_block(doc, text):
    """Agrega un bloque de cita con formato"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Limpiar el texto de asteriscos
    clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
    
    run = para.add_run(clean_text)
    run.font.italic = True
    run.font.size = Pt(10)
    
    # Agregar borde izquierdo
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement('w:left')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '24')
    border.set(qn('w:space'), '4')
    border.set(qn('w:color'), '4472C4')
    pBdr.append(border)
    pPr.append(pBdr)

def parse_table(lines, start_idx):
    """Parsea una tabla Markdown"""
    table_lines = []
    idx = start_idx
    while idx < len(lines) and '|' in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1
    
    if len(table_lines) < 2:
        return None, start_idx
    
    # Parsear filas
    rows = []
    for line in table_lines:
        if '---' in line and '|' in line:
            continue
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)
    
    return rows, idx

def add_table(doc, rows):
    """Agrega una tabla al documento"""
    if not rows or len(rows) < 1:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                # Limpiar formato Markdown
                clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
                cell.text = clean_text
                
                # Formato de encabezado (primera fila)
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                        # Sombreado de encabezado
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), 'E7E6E6')
                        cell._tc.get_or_add_tcPr().append(shd)
    
    doc.add_paragraph()  # Espacio despues de tabla

def process_inline_formatting(para, text):
    """Procesa formato inline (negrita, cursiva, codigo)"""
    # Patron para detectar diferentes formatos
    pattern = r'(\*\*\*([^*]+)\*\*\*|\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`|([^*`]+))'
    
    for match in re.finditer(pattern, text):
        full_match = match.group(0)
        
        if full_match.startswith('***') and full_match.endswith('***'):
            # Negrita + cursiva
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif full_match.startswith('**') and full_match.endswith('**'):
            # Negrita
            content = match.group(3)
            run = para.add_run(content)
            run.bold = True
        elif full_match.startswith('*') and full_match.endswith('*') and not full_match.startswith('**'):
            # Cursiva
            run = para.add_run(match.group(4))
            run.italic = True
        elif full_match.startswith('`') and full_match.endswith('`'):
            # Codigo inline
            run = para.add_run(match.group(5))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # Texto normal
            run = para.add_run(match.group(6) if match.group(6) else full_match)

def convert_md_to_docx(md_path, docx_path):
    """Convierte un archivo Markdown a Word"""
    print(f"Leyendo {md_path}...")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    doc = Document()
    set_document_styles(doc)
    
    # Configurar margenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    i = 0
    in_code_block = False
    code_block_content = []
    in_quote = False
    quote_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # Bloques de codigo
        if line.strip().startswith('```'):
            if in_code_block:
                # Fin del bloque
                add_code_block(doc, '\n'.join(code_block_content))
                code_block_content = []
                in_code_block = False
            else:
                # Inicio del bloque
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Citas (blockquotes)
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            add_quote_block(doc, quote_text)
            i += 1
            continue
        
        # Tablas
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            rows, new_idx = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
                i = new_idx
                continue
        
        # Lineas horizontales
        if line.strip() == '---' or line.strip() == '***':
            # Agregar linea horizontal
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Encabezados
        if line.startswith('######'):
            para = doc.add_paragraph(line[6:].strip(), style='Heading 4')
            i += 1
            continue
        elif line.startswith('#####'):
            para = doc.add_paragraph(line[5:].strip(), style='Heading 4')
            i += 1
            continue
        elif line.startswith('####'):
            para = doc.add_paragraph(line[4:].strip(), style='Heading 4')
            i += 1
            continue
        elif line.startswith('###'):
            para = doc.add_paragraph(line[3:].strip(), style='Heading 3')
            i += 1
            continue
        elif line.startswith('##'):
            para = doc.add_paragraph(line[2:].strip(), style='Heading 2')
            i += 1
            continue
        elif line.startswith('#'):
            para = doc.add_paragraph(line[1:].strip(), style='Heading 1')
            i += 1
            continue
        
        # Listas
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            list_text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(para, list_text)
            i += 1
            continue
        
        # Listas numeradas
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if numbered_match:
            list_text = numbered_match.group(2)
            para = doc.add_paragraph(style='List Number')
            process_inline_formatting(para, list_text)
            i += 1
            continue
        
        # Parrafos normales
        if line.strip():
            para = doc.add_paragraph()
            process_inline_formatting(para, line)
        
        i += 1
    
    # Guardar
    print(f"Guardando {docx_path}...")
    doc.save(docx_path)
    print("Conversion completada!")

if __name__ == '__main__':
    md_file = Path('INFORME_RECURSO_RD_CANNABIS.md')
    docx_file = Path('INFORME_RECURSO_RD_CANNABIS.docx')
    
    convert_md_to_docx(md_file, docx_file)
